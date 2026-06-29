"""設計レビュー資料のWord(.docx)を生成する。本文フォント: Meiryo UI。図はimages/のPNGを埋め込む。"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(__file__)
IMG = os.path.join(HERE, "images")
FONT = "Meiryo UI"
GREEN = RGBColor(0x1F, 0x6F, 0x54)


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def h(doc, text, level):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True,
             color=GREEN if level <= 2 else None)
    return p


def para(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def figure(doc, filename, caption, width=6.0):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_font(r, size=9, color=RGBColor(0x55, 0x55, 0x55))


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(hd)
        set_font(r, size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(val)
            set_font(r, size=9.5)
    return t


doc = Document()
style_base(doc)

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("医療機器アフターメンテナンス支援エージェント")
set_font(r, size=20, bold=True, color=GREEN)
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("設計レビュー資料（ピアレビュー用）")
set_font(r, size=13, bold=True)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("対象機種（合成）：内視鏡システム EVS-X1000（架空）／ "
                 "ベース：DataRobot Agentic Starter（LangGraph）")
set_font(r, size=9.5, color=RGBColor(0x55, 0x55, 0x55))

# 1
h(doc, "1. このアプリは何か（目的）", 1)
para(doc, "医療機器のアフターメンテナンス業務において、現場のフィールドサポートエンジニア（FSE）と"
          "HQのリモートサポートエンジニア（RSE）の判断・作業・引き継ぎを、Agentic RAG（マルチツール）で"
          "支援するモックアップである。狙いは次の2点。")
bullet(doc, "技能継承・現場支援：ベテランの暗黙知やマニュアル・修理履歴を横断検索し、根拠付きで提示する。")
bullet(doc, "業務プロセスの一気通貫支援：受電 → RSEトリアージ → FSE派遣 → 現地対応 → 報告書まで、"
            "1つのディスパッチ番号で串刺しして支援する。")
h(doc, "基本設計思想（最重要）", 3)
table(doc, ["#", "原則", "内容"], [
    ["0", "安全表示は不可侵", "安全に直結する注意事項（!）は⚠️黄色で必ず強調。削除・無効化しない。回帰テストで保護。"],
    ["1", "出力はドラフト、確定は人間", "生成物は叩き台。FSE/RSEが編集・追記・削除して最終確定（human-in-the-loop）。"],
    ["2", "全成果物が編集可能（同一UX）", "ヒアリング・トリアージ・作業手順・引き継ぎ・ブリーフィング・報告書を画面上で編集可能。"],
    ["3", "ペルソナは従業員IDで判別", "RSE####/FSE#### を従業員マスタ照合で検証してペルソナ確定。"],
    ["4", "運用ルールは差し替え可能に分離", "ディスパッチ発番ルール等はメーカー差異を dispatch_policy.py に集約。"],
])

# 2
h(doc, "2. 対象ペルソナ", 1)
bullet(doc, "RSE：受電トリアージ（電話クローズ or 派遣判断）、現地FSEの技術支援。最も知識・経験のあるシニア。")
bullet(doc, "FSE：現地で診断・修理。必要に応じてRSEに相談（エスカレーション）。")

# 3
h(doc, "3. 全体業務ライフサイクル", 1)
para(doc, "コールセンターが発番した1つのディスパッチ番号の配下に、RSE・FSEの全活動が紐づく。")
figure(doc, "01_lifecycle.png", "図1. 全体業務ライフサイクル", width=4.2)

# 4
h(doc, "4. RSE 局面①：顧客電話トリアージ", 1)
para(doc, "受電したRSEが「電話でクローズできるか／FSEを派遣するか」を判断する。保守契約の有無で有償/無償が分岐。")
figure(doc, "02_rse_phase1.png", "図2. RSE 局面①（顧客電話トリアージ）", width=4.6)
bullet(doc, "ヒアリング項目はエラーコード（カテゴリ）別の標準テンプレートで一貫化。既知情報は再確認しない。")
bullet(doc, "ヒアリングは状態・事実の確認に限定（お客様に部品交換等の作業を依頼しない）。")
bullet(doc, "結論は固定フォーマットのトリアージ表に集約（全項目RSEが編集可能）。")

# 5
h(doc, "5. RSE 局面②：現地FSEの技術支援（バックアップ）", 1)
para(doc, "既に派遣されたFSEが行き詰まって相談する局面。親ディスパッチ番号の配下に子番号（例 D-…-01）を採番。"
          "「使用中止・患者安全確保」等の顧客向け初動指示は出さず、技術支援に徹する。")

# 6
h(doc, "6. FSE 現地フロー", 1)
para(doc, "派遣で現場入りしたFSEは、RSEの引き継ぎ（ブリーフィング）を出発点として現地切り分けに合流する"
          "（症状の再入力は不要）。")
figure(doc, "03_fse_flow.png", "図3. FSE 現地フロー", width=4.6)
bullet(doc, "作業手順を出すメッセージには必ず結果確認の分岐を伴わせ、現地状況に応じて手順が進む。")
bullet(doc, "選択肢に無い観察は「その他（自由記述）」で受け、診断を見直す。")
bullet(doc, "完了時のSR報告書は顧客提出版（ですます調・有償部品の有無を明示・社内根拠は載せない）。")

# 7
h(doc, "7. ディスパッチ番号による串刺し（1案件＝1番号）", 1)
figure(doc, "04_dispatch.png", "図4. ディスパッチ番号による串刺し", width=6.0)
bullet(doc, "ヒアリング結果・トリアージ・ブリーフィング・作業履歴・報告書はすべて番号配下に保存（後日トラック可）。")
bullet(doc, "交代要員も「自分のIDで新チャット＋同じ番号」で過去経緯を引き継げる。")

# 8
h(doc, "8. データ／マスタ構成（合成データ）", 1)
figure(doc, "05_data.png", "図5. データ／マスタ構成", width=6.2)

# 9
h(doc, "9. エージェント構成とツール", 1)
para(doc, "LangGraph の単一エージェントが、以下のツール群（LangChain @tool）を状況に応じて呼び出す。")
figure(doc, "06_tools.png", "図6. エージェントとツールマップ", width=6.2)

# 10
h(doc, "10. UI上の主要コンポーネント", 1)
table(doc, ["マーカー", "UI", "役割"], [
    ["[[choices]]", "水色の選択肢ボタン（＋その他）", "分岐質問への回答"],
    ["[[hearing]]", "チェック付き表（メモ編集・行追加）", "顧客ヒアリング"],
    ["[[steps]]", "チェック付き作業表（⚠️安全強調）", "現地作業手順"],
    ["[[triage]]", "編集可能な表", "推定原因・類似傾向"],
    ["[[handoff_draft]]", "編集フォーム＋発行", "FSE→RSE引き継ぎ要約"],
    ["[[rse_actions]]", "編集表＋リリース", "RSE→FSEネクストアクション"],
    ["[[dispatch_briefing]]", "編集フォーム＋リリース", "HQ→営業所 派遣ブリーフィング"],
    ["[[report]]", "編集／プレビュー＋Word出力", "SR報告書ドラフト"],
    ["[[sources]]", "箇条書き＋モーダル（図面表示可）", "参照情報源"],
    ["[[complete_action]]", "全完了時の確認", "報告書ドラフト要否"],
])
bullet(doc, "画面左にプロセスマップを表示し、各ステップカードのクリックで該当箇所へジャンプ。")
bullet(doc, "アクションボタンは緑＋黒文字、分岐選択肢は水色＋黒文字で統一。")

# 11
h(doc, "11. 技術スタック", 1)
bullet(doc, "ベース：DataRobot Agentic Starter Application Template")
bullet(doc, "エージェント：LangGraph（datarobot_agent_class_from_langgraph + get_llm()、@tool）")
bullet(doc, "バックエンド：FastAPI（AG-UI ストリーミング）")
bullet(doc, "フロントエンド：React / Vite（チャットUI＋カスタムUIコンポーネント）")
bullet(doc, "LLM：DataRobot LLM Gateway 経由")

# 12
h(doc, "12. モックの制約・将来拡張", 1)
para(doc, "モックの割り切り（本実装では未対応）：", bold=True)
bullet(doc, "認証・アクセス制御・マルチテナント、部品在庫システム連携、院内ネットワーク制約対応。")
bullet(doc, "専用のASR（音声認識）・CV（画像認識）モデル。")
para(doc, "将来拡張ポイント：", bold=True)
bullet(doc, "参照情報源モーダルの図面に加え、動画・PDF・文書管理システム（DMS）連携。")
bullet(doc, "「その他（自由記述）」での音声・画像・動画入力（マルチモーダル）。")
bullet(doc, "ディスパッチ発番・担当アサイン・引き継ぎルールのメーカー別設定化。")
bullet(doc, "顧客マスタ／従業員マスタの認証基盤・基幹システム連携。")

# 13
h(doc, "13. ピアレビューでの確認観点（提案）", 1)
for i, t in enumerate([
    "業務フローの妥当性：RSE局面①/②、FSE現地、エスカレーション、有償/無償判断は実運用と整合するか。",
    "安全設計：安全注意（⚠️）の強調・不可侵化は十分か。患者・作業員安全の観点で漏れはないか。",
    "human-in-the-loop：全成果物の編集可能化と「確定は人間」の徹底は適切か。",
    "ペルソナ／権限：従業員ID検証・ペルソナ分岐・読み手の取り違え防止は妥当か。",
    "データ／引き継ぎ：ディスパッチ番号による串刺し・履歴保存・交代要員の引き継ぎは実用的か。",
    "拡張性：発番ポリシー分離・マスタ連携の差し替え可能性は将来要件に耐えるか。",
], 1):
    p = doc.add_paragraph(style="List Number")
    rr = p.add_run(t)
    set_font(rr, size=10.5)

out = os.path.join(HERE, "design-review.docx")
doc.save(out)
print("saved", out)
